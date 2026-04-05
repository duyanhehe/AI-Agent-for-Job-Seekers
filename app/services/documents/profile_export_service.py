from docx import Document
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from fastapi import HTTPException
import io

from app.models.cv_documents import CVDocuments
from app.models.user_profiles import UserProfile


def export_profile_docx_service(user_id: int, db: Session):
    """Build a DOCX from the user's primary CV profile and return a streaming response."""
    # GET PRIMARY CV
    cv = (
        db.query(CVDocuments)
        .filter(
            CVDocuments.user_id == user_id,
            CVDocuments.is_primary.is_(True),
        )
        .first()
    )

    if not cv:
        raise HTTPException(404, "Primary CV not found")

    profile_row = (
        db.query(UserProfile)
        .filter(
            UserProfile.user_id == user_id,
            UserProfile.cv_id == cv.id,
        )
        .first()
    )

    if not profile_row:
        raise HTTPException(404, "Profile not found")

    p = profile_row.profile

    # ---- BUILD DOC ----
    doc = Document()

    doc.add_heading(p.get("name", "Unnamed"), 0)

    contact = f"{p.get('email', '')} | {p.get('phone', '')} | {p.get('location', '')}"
    doc.add_paragraph(contact)

    # Education
    doc.add_heading("Education", level=1)
    for edu in p.get("education", []):
        doc.add_paragraph(
            f"{edu.get('school', '')} - {edu.get('degree', '')} ({edu.get('year', '')})"
        )

    # Experience
    doc.add_heading("Experience", level=1)
    for exp in p.get("work_experience", []):
        para = doc.add_paragraph()
        para.add_run(f"{exp.get('role', '')} - {exp.get('company', '')}\n").bold = True
        para.add_run(f"{exp.get('duration', '')}\n")
        para.add_run(exp.get("description", ""))

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(p.get("skills", [])))

    # ---- STREAM ----
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=resume.docx"},
    )
